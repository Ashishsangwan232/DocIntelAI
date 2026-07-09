"""
docx_loader.py
==============
DOCX text extraction using python-docx.

Word documents have no reliable notion of a "page" at the file-format
level (pagination is computed at render time by Word itself), so the
entire document is treated as a single logical unit with
`page_number=None`. Paragraphs and table cell text are both extracted
so tabular content in contracts/reports isn't silently dropped.
"""

from __future__ import annotations

from pathlib import Path

import docx
from docx.opc.exceptions import PackageNotFoundError

from src.loaders.base_loader import BaseLoader, LoadedDocument, PageContent
from src.utils.exceptions import CorruptedFileError
from src.utils.logger import get_logger

logger = get_logger(__name__)


class DocxLoader(BaseLoader):
    """Extracts paragraph and table text from DOCX files."""

    supported_extensions = ("docx",)

    def load(self, file_path: Path) -> LoadedDocument:
        try:
            document = docx.Document(str(file_path))
        except PackageNotFoundError as exc:
            logger.error("Failed to open DOCX %s: %s", file_path.name, exc)
            raise CorruptedFileError(
                f"'{file_path.name}' is not a valid DOCX file or is corrupted."
            ) from exc
        except Exception as exc:
            logger.error("Unexpected error opening DOCX %s: %s", file_path.name, exc)
            raise CorruptedFileError(
                f"Could not open '{file_path.name}': {exc}"
            ) from exc

        try:
            parts: list[str] = [
                p.text for p in document.paragraphs if p.text and p.text.strip()
            ]
            for table in document.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        parts.append(row_text)
        except Exception as exc:
            logger.error("Failed to extract text from DOCX %s: %s", file_path.name, exc)
            raise CorruptedFileError(
                f"Text extraction failed for '{file_path.name}': {exc}"
            ) from exc

        full_text = "\n".join(parts)
        logger.info(
            "Loaded DOCX '%s' with %d paragraphs/rows", file_path.name, len(parts)
        )
        return LoadedDocument(
            source_filename=file_path.name,
            pages=[PageContent(text=full_text, page_number=None)],
        )
