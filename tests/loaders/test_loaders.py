"""
Unit tests for src/loaders/*.

Uses real PDF and DOCX files generated on the fly with PyMuPDF and
python-docx (not mocks) so extraction logic is verified against the
actual file formats it will encounter in production.
"""

from __future__ import annotations

from pathlib import Path

import docx
import fitz
import pytest

from src.loaders.docx_loader import DocxLoader
from src.loaders.loader_factory import get_loader, supported_extensions
from src.loaders.pdf_loader import PDFLoader
from src.loaders.txt_loader import TextLoader
from src.utils.exceptions import CorruptedFileError, UnsupportedFileTypeError


@pytest.fixture()
def sample_pdf(tmp_path: Path) -> Path:
    path = tmp_path / "sample.pdf"
    pdf = fitz.open()
    page1 = pdf.new_page()
    page1.insert_text((72, 72), "Page one content about warranties.")
    page2 = pdf.new_page()
    page2.insert_text((72, 72), "Page two content about termination.")
    pdf.save(str(path))
    pdf.close()
    return path


@pytest.fixture()
def corrupted_pdf(tmp_path: Path) -> Path:
    path = tmp_path / "corrupted.pdf"
    path.write_bytes(b"this is not a valid pdf file")
    return path


@pytest.fixture()
def sample_docx(tmp_path: Path) -> Path:
    path = tmp_path / "sample.docx"
    document = docx.Document()
    document.add_paragraph("This paragraph discusses liability limits.")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Clause"
    table.rows[0].cells[1].text = "Liability"
    document.save(str(path))
    return path


class TestPDFLoader:
    def test_extracts_text_per_page(self, sample_pdf: Path) -> None:
        loaded = PDFLoader().load(sample_pdf)
        assert loaded.page_count == 2
        assert "warranties" in loaded.pages[0].text
        assert "termination" in loaded.pages[1].text
        assert loaded.pages[0].page_number == 1
        assert loaded.pages[1].page_number == 2

    def test_corrupted_pdf_raises(self, corrupted_pdf: Path) -> None:
        with pytest.raises(CorruptedFileError):
            PDFLoader().load(corrupted_pdf)


class TestDocxLoader:
    def test_extracts_paragraphs_and_tables(self, sample_docx: Path) -> None:
        loaded = DocxLoader().load(sample_docx)
        assert loaded.page_count is None
        assert "liability limits" in loaded.full_text
        assert "Clause" in loaded.full_text and "Liability" in loaded.full_text

    def test_corrupted_docx_raises(self, tmp_path: Path) -> None:
        bad_path = tmp_path / "bad.docx"
        bad_path.write_bytes(b"not a docx")
        with pytest.raises(CorruptedFileError):
            DocxLoader().load(bad_path)


class TestTextLoader:
    def test_loads_txt_file(self, tmp_path: Path) -> None:
        path = tmp_path / "notes.txt"
        path.write_text("Plain text notes about the project.")
        loaded = TextLoader().load(path)
        assert loaded.page_count is None
        assert "Plain text notes" in loaded.full_text

    def test_loads_markdown_file(self, tmp_path: Path) -> None:
        path = tmp_path / "readme.md"
        path.write_text("# Title\n\nSome markdown body text.")
        loaded = TextLoader().load(path)
        assert "markdown body text" in loaded.full_text

    def test_falls_back_to_latin1_encoding(self, tmp_path: Path) -> None:
        path = tmp_path / "latin.txt"
        path.write_bytes("café".encode("latin-1"))
        loaded = TextLoader().load(path)
        assert loaded.full_text  # decodes without raising


class TestLoaderFactory:
    def test_returns_pdf_loader(self) -> None:
        assert isinstance(get_loader("pdf"), PDFLoader)

    def test_returns_docx_loader(self) -> None:
        assert isinstance(get_loader("docx"), DocxLoader)

    def test_txt_and_md_share_text_loader(self) -> None:
        assert isinstance(get_loader("txt"), TextLoader)
        assert isinstance(get_loader("md"), TextLoader)

    def test_case_insensitive_lookup(self) -> None:
        assert isinstance(get_loader("PDF"), PDFLoader)

    def test_unsupported_extension_raises(self) -> None:
        with pytest.raises(UnsupportedFileTypeError):
            get_loader("exe")

    def test_supported_extensions_includes_all_formats(self) -> None:
        exts = supported_extensions()
        assert {"pdf", "docx", "txt", "md"}.issubset(set(exts))
